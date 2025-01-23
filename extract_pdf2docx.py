import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"正在安装 {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()[package] = __import__(package)

# 检查并安装所需的库
install_and_import('pdfplumber')
install_and_import('docx')  # 对应于 python-docx
install_and_import('re')
install_and_import('tqdm')

import pdfplumber
import re
from docx import Document
import argparse
import os
from tqdm import tqdm

def is_title(text):
    # 使用正则表达式判断是否为标题
    pattern = re.compile(r"^[0-9]+\.[0-9]+.*")
    return bool(pattern.match(text.strip()))

def merge_paragraphs(paragraphs):
    merged_text = []
    i = 0

    while i < len(paragraphs):
        para = paragraphs[i].strip()
        if not para or is_title(para):
            merged_text.append(para)
            i += 1
            continue

        # 合并段落
        while i + 1 < len(paragraphs):
            next_para = paragraphs[i + 1].strip()
            if is_title(next_para):
                break

            # 检查段尾符号
            if para[-1] not in ".!?:":
                para += " " + next_para
                i += 1
            else:
                break

        merged_text.append(para)
        i += 1

    return merged_text

def extract_text_with_pdfplumber(pdf_path, skip_header=False, skip_footer=False, mixed_footnotes=False):
    body_text = []
    footnotes = []
    current_footnote = None
    last_was_number_footnote = False

    with pdfplumber.open(pdf_path) as pdf:
        for page in tqdm(pdf.pages, desc="Processing pages"):
            lines = page.extract_text().split('\n')
            
            # 跳过页眉
            if skip_header and lines:
                lines = lines[1:]
            
            # 跳过页脚
            if skip_footer and lines:
                lines = lines[:-1]
            
            for line in lines:
                # 检查行是否以数字或字母开头，并且后面不是标题格式或年份
                if line and (line[0].isdigit() or (mixed_footnotes and last_was_number_footnote and line[0].isalpha())):
                    # 使用正则表达式检查是否是标题格式
                    if re.match(r'^\d+(\.\d+)*\.\s*\w', line):
                        body_text.append(line)
                        current_footnote = None  # 结束当前脚注
                        last_was_number_footnote = False
                    # 检查是否是年份（假设年份为4位数字）
                    elif len(line) >= 4 and line[:4].isdigit() and (line[4].isspace() or line[4] in {'.', ',', ';', ':'}):
                        body_text.append(line)
                        current_footnote = None  # 结束当前脚注
                        last_was_number_footnote = False
                    # 检查是否是脚注
                    elif re.match(r'^\d+[a-zA-Z]?[\.\s]', line):
                        if current_footnote is not None:
                            footnotes.append(current_footnote)
                        current_footnote = line
                        last_was_number_footnote = line[0].isdigit()
                    elif mixed_footnotes and last_was_number_footnote and re.match(r'^[a-zA-Z][\.\s]', line):
                        # 处理字母脚注，只在数字脚注后面
                        if current_footnote is not None:
                            footnotes.append(current_footnote)
                        current_footnote = line
                        last_was_number_footnote = False
                    else:
                        # 如果当前正在记录脚注，继续添加行
                        if current_footnote is not None:
                            current_footnote += " " + line
                        else:
                            body_text.append(line)
                else:
                    if current_footnote is not None:
                        # 如果当前正在记录脚注，继续添加行
                        current_footnote += " " + line
                    else:
                        body_text.append(line)

            # 如果页面结束时还有未结束的脚注，添加到脚注列表
            if current_footnote is not None:
                footnotes.append(current_footnote)
                current_footnote = None

    # 合并段落
    body_text = merge_paragraphs(body_text)

    return body_text, footnotes

def save_to_docx(body_text, footnotes, output_file):
    doc = Document()
    doc.add_heading('正文', level=1)
    for paragraph in body_text:
        doc.add_paragraph(paragraph)
    
    doc.add_page_break()
    doc.add_heading('脚注', level=1)
    for footnote in footnotes:
        doc.add_paragraph(footnote)
    
    doc.save(output_file)

def process_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        full_text = ""
        for page in pdf.pages:
            full_text += page.extract_text() + "\n"

    merged_text = merge_paragraphs(full_text)
    return merged_text

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from PDF and save to DOCX.")
    parser.add_argument("pdf_path", nargs='?', help="Path to the input PDF file.")
    parser.add_argument("--skip-header", action="store_true", help="Skip the header line on each page.")
    parser.add_argument("--skip-footer", action="store_true", help="Skip the footer line on each page.")
    parser.add_argument("--mixed-footnotes", action="store_true", help="Indicate if there are letter footnotes following number footnotes.")
    
    args = parser.parse_args()

    if not args.pdf_path:
        parser.print_help()
        sys.exit(1)
    
    # 设置输出文件名与输入文件同名，但扩展名为.docx
    output_file = os.path.splitext(args.pdf_path)[0] + ".docx"
    
    body_text, footnotes = extract_text_with_pdfplumber(
        args.pdf_path, 
        skip_header=args.skip_header, 
        skip_footer=args.skip_footer,
        mixed_footnotes=args.mixed_footnotes
    )
    save_to_docx(body_text, footnotes, output_file)